#! /usr/bin/env python3
# wordPrac.py - write a Resume-template generator

import docx
from docx.shared import Pt

doc = docx.Document()

def greeting():
	print('The following program will generate your provided' +
		' information in a professional resume template.' +
		' Please enter all information as you\'d like it to'+
		' appear.')

	nameRole()

def nameRole():
	
	print('')
	print('Let\'s get started. What\'s your professional name?')
	
	name = input()
	nameHead = doc.add_paragraph(name, 'Title')
	nameHead.alignment = 0
	nameHead.runs[0].font.size = Pt(25)

	print('What\'s your general role? ie Web Developer, '+
		'Writer, Copyeditor, etc.')

	role = input()
	nameHead.add_run(' '+role).italic = True
	doc.paragraphs[0].runs[1].style = "SubtitleChar"
	
	print('')
	print('Next, we\'ll handle contact info.')
	print('What\'s your email?')
	email = input()
	contactSect = doc.add_paragraph(email, 'NoSpacing')
	contactSect.alignment = 1

	print('What\'s your phone number?')
	phone = input()

	contactSect.add_run('  |  ' + phone)


	summary()

def summary():
	print('Write a quick blurb about yourself. What are your '+
		'strengths, what are you looking to learn?')

	summary = input()

	doc.add_heading('Summary:', 2)
	summaryInfo = doc.add_paragraph()
	summaryInfo.add_run(summary)
	summaryInfo.runs[0].style = 'SubtleEmphasis'
	
	workExp()

def workExp():

	
	
	
	

	
	while True:

		print('')

		print('Type \'Work\' to add Work Experience. Type '+
			'\'Extra\' to add Extracurriculars. ' +
			'Type \'done\' if finished entering both '+
			'work and extracurricular info.')

		nextSect = input()

		if nextSect == 'Work':
			workHead = doc.add_heading('Work Experience:', 2)
		elif nextSect == 'Extra':
			workHead = doc.add_heading('Extracurriculars:', 2)
		elif nextSect == 'done':
			break

		while True:
			print('Add your next most recent item name.')
			lastJob = input()
			workHead.add_run(lastJob)
			workHead.style = 'NoSpacing'
			workHead.runs[0].bold = True

			print('What were the dates worked/participated?')
			lastDates = input()
			workHead.add_run('\t\t\t\t\t' +lastDates).italic = True

			print('What was your title?')
			lastTitle = input()
			workHead.runs[1].add_break()
			workHead.add_run(lastTitle)
			workHead.style = 'NoSpacing'

			print('What skills did you learn? Enter one at a time '+
				'followed by \'Enter\', then type done to move on.')
			skills = []
			
			while True:
				skill = input()
				if skill == 'done':
					break
				else:
					skills.append(skill)
					continue

			for each in skills: 
				listed = doc.add_paragraph(each, style='ListBullet2')
			
			print('To add to the same section, type \'add\', '+
				'followed by the enter key. Otherwise, click \'done\'')

			answer = input()
			if answer == 'add':
				continue
			elif answer == 'done':
				break

#	education()

def education():
	doc.add_heading('Education:', 2)
	print('The following will ask about your education.')

	while True:
		print('')
		print('Please proceed with the next most recent educational experience: ')
		print('What\'s the name of the institution? Type \'done\' if finished '+
				'entering education information.')
		edu = input()
		if edu == 'done':
			break
		else: 
			eduHead = doc.add_paragraph()
			eduHead.add_run(edu)
			eduHead.style = 'Normal'
			eduHead.runs[0].bold = True

			print('What are the years attended?')
			dates = input()

			eduHead.add_run('\t\t\t\t\t\t\t\t' + dates).italic = True

			print('Was it a B.A. or B.S., or something else? If '+
				'something else, please '+
				'enter the name of the degree.')
			deg = input()
			eduHead.runs[1].add_break()
			eduHead.add_run(deg)
			
			print('What was the degree in?')
			degIn = input()
			eduHead.add_run(', ' + degIn)
			
			print('If you graduated with any honors (cum laude, '+
				'magna cum laude, etc.), list them here. If not, hit enter.')
			honors = input()
			if honors == '':
				eduHead.add_run('')
			else:
				eduHead.add_run('\t\t\t\t\t\t\t\t\t' + '[' + honors + 
					']').italic = True
			
			print('What was your GPA? If not applicable, just hit enter.')
			gpa = input()
			eduHead.runs[4].add_break()
			if gpa == '':
				eduHead.add_run()
			else:
				eduHead.add_run('GPA: ' + gpa)



#greeting()
#nameRole()
workExp()
#education()
#contact()



doc.save('resumeTemplate.docx')